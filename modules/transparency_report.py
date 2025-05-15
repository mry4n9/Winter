from docx import Document
from docx.shared import Pt
from io import BytesIO

def create_report_docx(company_name: str, 
                       url_context_raw: str, url_context_sum: str,
                       additional_context_raw: str, additional_context_sum: str,
                       lead_magnet_raw: str, lead_magnet_sum: str) -> BytesIO:
    """
    Creates a DOCX transparency report and returns it as a BytesIO stream.
    """
    doc = Document()

    # Title
    title = doc.add_heading("AI Transparency Report", level=1)
    # Note: Default alignment is left. No specific font styling requested beyond default.

    doc.add_paragraph(
        f"This document outlines the source text and AI-generated summaries used for creating ad copy for {company_name}."
    )
    doc.add_paragraph()  # Spacer

    # URL Context
    doc.add_heading("Raw Extract from URL:", level=2)
    doc.add_paragraph(url_context_raw if url_context_raw else "N/A (No URL provided or extraction failed)")
    doc.add_paragraph()
    doc.add_heading("AI Summarized (URL):", level=2)
    doc.add_paragraph(url_context_sum if url_context_sum else "N/A (No summary generated)")
    doc.add_paragraph()

    # Additional Upload Context
    doc.add_heading("Extract from Additional Upload:", level=2)
    doc.add_paragraph(additional_context_raw if additional_context_raw else "N/A (No additional context file uploaded or extraction failed)")
    doc.add_paragraph()
    doc.add_heading("AI Summarized (Additional Upload):", level=2)
    doc.add_paragraph(additional_context_sum if additional_context_sum else "N/A (No summary generated)")
    doc.add_paragraph()

    # Lead Magnet Context
    doc.add_heading("Extract from Lead Magnet:", level=2)
    doc.add_paragraph(lead_magnet_raw if lead_magnet_raw else "N/A (No lead magnet file uploaded or extraction failed)")
    doc.add_paragraph()
    doc.add_heading("AI Summarized (Lead Magnet):", level=2)
    doc.add_paragraph(lead_magnet_sum if lead_magnet_sum else "N/A (No summary generated)")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream